"""
Patched OutlookMessageLoader that handles emails with no plain text body.
This fixes the "page_content Input should be a valid string" error.
Enhanced to properly extract links, attachments, and structured content.
Supports: attachment content extraction, nested emails, QR code URL extraction.
"""
import os
import re
import io
import tempfile
from pathlib import Path
from typing import Iterator, Union, List, Tuple, Dict, Any, Optional
from langchain_core.documents import Document
from langchain_community.document_loaders.base import BaseLoader


def extract_links_from_html(html_content: str) -> List[Tuple[str, str]]:
    """Extract links from HTML content as (text, url) tuples."""
    links = []
    # Match <a> tags with href
    pattern = r'<a\s+[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>'
    for match in re.finditer(pattern, html_content, re.IGNORECASE | re.DOTALL):
        url = match.group(1)
        text = re.sub(r'<[^>]+>', '', match.group(2)).strip()  # Remove nested tags
        if url and not url.startswith('#') and not url.startswith('mailto:'):
            links.append((text or url, url))
        elif url and url.startswith('mailto:'):
            # Extract email from mailto links
            email = url.replace('mailto:', '').split('?')[0]
            links.append((text or email, email))
    return links


def html_to_text(html_content: str) -> str:
    """Convert HTML to readable text, preserving structure and links."""
    if not html_content:
        return ""

    text = html_content

    # Decode if bytes
    if isinstance(text, bytes):
        text = text.decode('utf-8', errors='replace')

    # Remove style and script blocks entirely
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.IGNORECASE | re.DOTALL)

    # Convert links to readable format: "text (url)" or just "url"
    def replace_link(match):
        url = match.group(1)
        link_text = re.sub(r'<[^>]+>', '', match.group(2)).strip()
        if url.startswith('mailto:'):
            email = url.replace('mailto:', '').split('?')[0]
            return f"{link_text} <{email}>" if link_text and link_text != email else f"<{email}>"
        if link_text and link_text != url and not link_text.startswith('http'):
            return f"{link_text} ({url})"
        return url

    text = re.sub(r'<a\s+[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
                  replace_link, text, flags=re.IGNORECASE | re.DOTALL)

    # Convert block elements to newlines
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</tr>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</td>', '\t', text, flags=re.IGNORECASE)
    text = re.sub(r'</li>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<h[1-6][^>]*>', '\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'</h[1-6]>', '\n', text, flags=re.IGNORECASE)

    # Convert list markers
    text = re.sub(r'<li[^>]*>', '\n• ', text, flags=re.IGNORECASE)

    # Remove remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)

    # Decode HTML entities
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&quot;', '"', text)
    text = re.sub(r'&#39;', "'", text)
    text = re.sub(r'&rsquo;', "'", text)
    text = re.sub(r'&lsquo;', "'", text)
    text = re.sub(r'&rdquo;', '"', text)
    text = re.sub(r'&ldquo;', '"', text)
    text = re.sub(r'&ndash;', '-', text)
    text = re.sub(r'&mdash;', '—', text)
    text = re.sub(r'&bull;', '•', text)
    text = re.sub(r'&#\d+;', '', text)  # Remove remaining numeric entities

    # Clean up whitespace
    text = re.sub(r'[ \t]+', ' ', text)  # Collapse horizontal whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Max 2 consecutive newlines
    text = '\n'.join(line.strip() for line in text.split('\n'))  # Strip each line
    text = text.strip()

    return text


def extract_qr_codes_from_image(image_data: bytes) -> List[str]:
    """Extract URLs from QR codes in an image. Returns list of decoded URLs."""
    urls = []
    try:
        from PIL import Image
        import pyzbar.pyzbar as pyzbar

        # Open image from bytes
        img = Image.open(io.BytesIO(image_data))

        # Convert to RGB if necessary (for grayscale or RGBA)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')

        # Decode QR codes
        decoded_objects = pyzbar.decode(img)

        for obj in decoded_objects:
            if obj.type == 'QRCODE':
                data = obj.data.decode('utf-8', errors='replace')
                # Check if it looks like a URL
                if data.startswith('http://') or data.startswith('https://') or data.startswith('www.'):
                    urls.append(data)
                elif re.match(r'^[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', data):
                    # Looks like a domain, add https://
                    urls.append(f'https://{data}')
                else:
                    # Non-URL QR data, still include it marked as QR data
                    urls.append(f'[QR Data: {data}]')
    except ImportError:
        # Try alternative: jsqr via subprocess (if available)
        try:
            import subprocess
            import json
            import base64

            # Write image to temp file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(image_data)
                tmp_path = tmp.name

            # Try Node.js-based QR detection if available
            script = f'''
            const Jimp = require('jimp');
            const jsQR = require('jsqr');

            (async () => {{
                const image = await Jimp.read('{tmp_path}');
                const code = jsQR(image.bitmap.data, image.bitmap.width, image.bitmap.height);
                if (code) {{
                    console.log(JSON.stringify({{data: code.data}}));
                }}
            }})();
            '''
            result = subprocess.run(['node', '-e', script], capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                data = json.loads(result.stdout.strip())
                if data.get('data'):
                    qr_data = data['data']
                    if qr_data.startswith('http'):
                        urls.append(qr_data)
                    else:
                        urls.append(f'[QR Data: {qr_data}]')

            os.unlink(tmp_path)
        except Exception:
            pass
    except Exception:
        pass

    return urls


def extract_text_from_pdf(pdf_data: bytes) -> str:
    """Extract text from PDF bytes."""
    text = ""
    try:
        # Try PyPDF2 first
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(pdf_data))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except ImportError:
            # Try pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                pass
    except Exception:
        pass
    return text.strip()


def extract_text_from_docx(docx_data: bytes) -> str:
    """Extract text from DOCX bytes."""
    text = ""
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_data))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)
    except ImportError:
        pass
    except Exception:
        pass
    return text.strip()


def extract_text_from_xlsx(xlsx_data: bytes) -> str:
    """Extract text from XLSX bytes."""
    text = ""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_data), data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"\n[Sheet: {sheet_name}]\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text += row_text + "\n"
    except ImportError:
        pass
    except Exception:
        pass
    return text.strip()


def parse_nested_email(email_data: bytes, file_ext: str = '.eml') -> Dict[str, Any]:
    """Parse a nested email attachment and return its content."""
    result = {
        'subject': '',
        'from': '',
        'to': '',
        'body': '',
        'attachments': []
    }

    try:
        if file_ext.lower() == '.msg':
            import extract_msg
            # Save to temp file for extract_msg
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as tmp:
                tmp.write(email_data)
                tmp_path = tmp.name

            msg = extract_msg.Message(tmp_path)
            result['subject'] = msg.subject or ''
            result['from'] = msg.sender or ''
            result['to'] = msg.to or ''
            result['body'] = msg.body or ''
            if not result['body'] and msg.htmlBody:
                result['body'] = html_to_text(msg.htmlBody)
            msg.close()
            os.unlink(tmp_path)
        else:
            # Standard .eml format
            import email
            from email.header import decode_header

            msg = email.message_from_bytes(email_data)

            # Decode subject
            subject = msg.get('Subject', '')
            if subject:
                decoded_parts = decode_header(subject)
                subject = ''
                for content, encoding in decoded_parts:
                    if isinstance(content, bytes):
                        subject += content.decode(encoding or 'utf-8', errors='replace')
                    else:
                        subject += content

            result['subject'] = subject
            result['from'] = msg.get('From', '')
            result['to'] = msg.get('To', '')

            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == 'text/plain' and not result['body']:
                        try:
                            result['body'] = part.get_payload(decode=True).decode('utf-8', errors='replace')
                        except:
                            pass
                    elif content_type == 'text/html' and not result['body']:
                        try:
                            html = part.get_payload(decode=True).decode('utf-8', errors='replace')
                            result['body'] = html_to_text(html)
                        except:
                            pass
            else:
                try:
                    result['body'] = msg.get_payload(decode=True).decode('utf-8', errors='replace')
                except:
                    pass
    except Exception:
        pass

    return result


def extract_attachment_content(att_data: bytes, filename: str, mimetype: str = None) -> Dict[str, Any]:
    """Extract content from an attachment based on its type."""
    result = {
        'type': 'unknown',
        'text': '',
        'qr_urls': [],
        'nested_email': None
    }

    if not att_data:
        return result

    ext = os.path.splitext(filename.lower())[1] if filename else ''

    # Image files - check for QR codes
    if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'] or (mimetype and mimetype.startswith('image/')):
        result['type'] = 'image'
        qr_urls = extract_qr_codes_from_image(att_data)
        if qr_urls:
            result['qr_urls'] = qr_urls

    # PDF files
    elif ext == '.pdf' or mimetype == 'application/pdf':
        result['type'] = 'pdf'
        result['text'] = extract_text_from_pdf(att_data)

    # Word documents
    elif ext in ['.docx'] or mimetype in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        result['type'] = 'docx'
        result['text'] = extract_text_from_docx(att_data)

    # Excel files
    elif ext in ['.xlsx'] or mimetype in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
        result['type'] = 'xlsx'
        result['text'] = extract_text_from_xlsx(att_data)

    # Plain text files
    elif ext in ['.txt', '.csv', '.log', '.json', '.xml', '.html', '.htm', '.md', '.py', '.js', '.ts', '.css']:
        result['type'] = 'text'
        try:
            result['text'] = att_data.decode('utf-8', errors='replace')
        except:
            result['text'] = str(att_data)

    # Nested email files
    elif ext in ['.eml', '.msg'] or mimetype in ['message/rfc822', 'application/vnd.ms-outlook']:
        result['type'] = 'email'
        result['nested_email'] = parse_nested_email(att_data, ext)

    return result


class PatchedOutlookMessageLoader(BaseLoader):
    """
    Patched Outlook Message loader that handles emails with no plain text body.
    Falls back to HTML body with proper link extraction.
    Includes attachments and structured metadata.
    Supports attachment content extraction, nested emails, and QR code detection.
    """

    def __init__(self, file_path: Union[str, Path], extract_attachment_content: bool = True):
        """Initialize with a file path.

        Args:
            file_path: The path to the Outlook Message file.
            extract_attachment_content: Whether to extract content from attachments.
        """
        self.file_path = str(file_path)
        self.extract_attachments = extract_attachment_content

        if not os.path.isfile(self.file_path):
            raise ValueError(f"File path {self.file_path} is not a valid file")

        try:
            import extract_msg  # noqa:F401
        except ImportError:
            raise ImportError(
                "extract_msg is not installed. Please install it with "
                "`pip install extract_msg`"
            )

    def lazy_load(self) -> Iterator[Document]:
        import extract_msg

        msg = extract_msg.Message(self.file_path)

        # Build content sections
        content_parts = []

        # Email header info
        header_parts = []
        if msg.subject:
            header_parts.append(f"Subject: {msg.subject}")
        if msg.sender:
            header_parts.append(f"From: {msg.sender}")
        if msg.to:
            header_parts.append(f"To: {msg.to}")
        if msg.cc:
            header_parts.append(f"CC: {msg.cc}")
        if msg.date:
            header_parts.append(f"Date: {msg.date}")

        if header_parts:
            content_parts.append("=== EMAIL HEADER ===")
            content_parts.extend(header_parts)
            content_parts.append("")

        # Email body
        body_content = None
        links_extracted = []

        # First try plain text body
        if msg.body and msg.body.strip():
            body_content = msg.body.strip()
            # Try to extract URLs from plain text too
            url_pattern = r'https?://[^\s<>"\')\]]+|www\.[^\s<>"\')\]]+'
            found_urls = re.findall(url_pattern, body_content)
            links_extracted = [(url, url) for url in found_urls]

        # Fall back to HTML body with proper parsing
        elif msg.htmlBody:
            html_content = msg.htmlBody
            if isinstance(html_content, bytes):
                html_content = html_content.decode('utf-8', errors='replace')

            # Extract links before converting to text
            links_extracted = extract_links_from_html(html_content)

            # Convert HTML to readable text
            body_content = html_to_text(html_content)

        if body_content:
            content_parts.append("=== EMAIL BODY ===")
            content_parts.append(body_content)
            content_parts.append("")

        # Extract and list links
        if links_extracted:
            content_parts.append("=== LINKS FOUND ===")
            seen_urls = set()
            for text, url in links_extracted:
                if url not in seen_urls:
                    seen_urls.add(url)
                    if text and text != url:
                        content_parts.append(f"• {text}: {url}")
                    else:
                        content_parts.append(f"• {url}")
            content_parts.append("")

        # Process attachments
        attachments_info = []
        attachment_contents = []
        nested_emails = []
        qr_urls_found = []

        try:
            if hasattr(msg, 'attachments') and msg.attachments:
                for att in msg.attachments:
                    att_name = getattr(att, 'longFilename', None) or getattr(att, 'shortFilename', None) or 'unnamed'
                    att_size = getattr(att, 'size', None)
                    att_data = getattr(att, 'data', None)
                    att_mimetype = getattr(att, 'mimetype', None)

                    # Format attachment info string
                    if att_size:
                        if att_size > 1024 * 1024:
                            size_str = f"{att_size / (1024*1024):.1f} MB"
                        elif att_size > 1024:
                            size_str = f"{att_size / 1024:.1f} KB"
                        else:
                            size_str = f"{att_size} bytes"
                        attachments_info.append(f"• {att_name} ({size_str})")
                    else:
                        attachments_info.append(f"• {att_name}")

                    # Extract content from attachment if enabled
                    if self.extract_attachments and att_data:
                        extracted = extract_attachment_content(att_data, att_name, att_mimetype)

                        if extracted['text']:
                            attachment_contents.append({
                                'filename': att_name,
                                'type': extracted['type'],
                                'text': extracted['text'][:5000]  # Limit to 5000 chars per attachment
                            })

                        if extracted['qr_urls']:
                            qr_urls_found.extend(extracted['qr_urls'])

                        if extracted['nested_email']:
                            nested_emails.append({
                                'filename': att_name,
                                **extracted['nested_email']
                            })
        except Exception:
            pass  # Ignore attachment extraction errors

        if attachments_info:
            content_parts.append("=== ATTACHMENTS ===")
            content_parts.extend(attachments_info)
            content_parts.append("")

        # Add QR code URLs if found
        if qr_urls_found:
            content_parts.append("=== QR CODE URLS (from image attachments) ===")
            for url in qr_urls_found:
                content_parts.append(f"• {url}")
            content_parts.append("")

        # Add attachment content extracts
        if attachment_contents:
            content_parts.append("=== ATTACHMENT CONTENT ===")
            for att_content in attachment_contents:
                content_parts.append(f"\n--- {att_content['filename']} ({att_content['type']}) ---")
                # Truncate long content
                text = att_content['text']
                if len(text) > 2000:
                    text = text[:2000] + "\n[... content truncated ...]"
                content_parts.append(text)
            content_parts.append("")

        # Add nested email content
        if nested_emails:
            content_parts.append("=== NESTED EMAILS ===")
            for nested in nested_emails:
                content_parts.append(f"\n--- Attached Email: {nested['filename']} ---")
                if nested.get('subject'):
                    content_parts.append(f"Subject: {nested['subject']}")
                if nested.get('from'):
                    content_parts.append(f"From: {nested['from']}")
                if nested.get('to'):
                    content_parts.append(f"To: {nested['to']}")
                if nested.get('body'):
                    body = nested['body']
                    if len(body) > 2000:
                        body = body[:2000] + "\n[... content truncated ...]"
                    content_parts.append(f"\n{body}")
            content_parts.append("")

        # Combine all content
        full_content = '\n'.join(content_parts).strip()

        # Ensure we have some content
        if not full_content:
            full_content = "(Empty email - no readable content found)"

        # Build metadata
        metadata = {
            "source": self.file_path,
            "subject": msg.subject or "",
            "sender": msg.sender or "",
            "date": str(msg.date) if msg.date else "",
            "type": "email",
            "format": "msg",
        }

        # Add recipients if available
        if msg.to:
            metadata["to"] = msg.to
        if msg.cc:
            metadata["cc"] = msg.cc

        # Add links to metadata for easy access
        if links_extracted:
            metadata["links"] = [url for _, url in links_extracted]
            metadata["link_count"] = len(links_extracted)

        # Add QR URLs to metadata
        if qr_urls_found:
            metadata["qr_urls"] = qr_urls_found
            metadata["qr_url_count"] = len(qr_urls_found)

        # Add attachment info to metadata
        if attachments_info:
            metadata["attachments"] = [a.lstrip('• ') for a in attachments_info]
            metadata["attachment_count"] = len(attachments_info)

        # Add nested email count
        if nested_emails:
            metadata["nested_email_count"] = len(nested_emails)

        msg.close()

        yield Document(
            page_content=full_content,
            metadata=metadata,
        )
